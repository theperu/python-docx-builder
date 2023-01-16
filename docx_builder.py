import json
from docx import Document

# Open JSON file and load data
with open('data.json') as json_file:
    data = json.load(json_file)

# Open template .docx file
document = Document('template.docx')

# replace placeholders with data from JSON
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        #print(run.text)
        for key, value in data.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

# Loop through tables in the document
for table in document.tables:
    # Loop through rows in the table
    for row in table.rows:
        # Loop through cells in the row
        for cell in row.cells:
            # Loop through paragraphs in the cell
            for paragraph in cell.paragraphs:
                # Loop through runs in the paragraph
                for run in paragraph.runs:
                    # Replace placeholders with data
                    for key, value in data.items():
                        if key in run.text:
                            run.text = run.text.replace(key, value)

# Loop through sections in the document
for section in document.sections:
    # Check if the section has a header
    if section.header:
        # Loop through paragraphs in the header
        for paragraph in section.header.paragraphs:
            # Loop through runs in the paragraph
            for run in paragraph.runs:
                # Replace placeholders with data
                for key, value in data.items():
                    run.text = run.text.replace(key, value)

# Save the modified document
document.save('filled_document.docx')